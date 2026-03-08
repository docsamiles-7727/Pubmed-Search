"""Output generation: Markdown (.md) and Word (.docx) files.

Converts the LLM-generated summary markdown into professionally formatted documents.
"""

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config


def _log(msg: str):
    sys.stderr.write(f"[Output] {msg}\n")
    sys.stderr.flush()


def _slugify(text: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    return re.sub(r"[\s_]+", "_", slug).strip("_")[:60]


def _ensure_output_dir(topic: str) -> Path:
    slug = _slugify(topic)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = config.OUTPUT_DIR / f"{slug}_{timestamp}"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def save_markdown(summary_md: str, topic: str, output_dir: Path = None) -> Path:
    """Save the summary as a markdown file."""
    if output_dir is None:
        output_dir = _ensure_output_dir(topic)
    md_path = output_dir / "literature_review.md"

    header = f"# Literature Review: {topic}\n\n"
    header += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n---\n\n"

    md_path.write_text(header + summary_md, encoding="utf-8")
    _log(f"Markdown saved to {md_path}")
    return md_path


def save_docx(summary_md: str, topic: str, output_dir: Path = None) -> Path:
    """Convert markdown summary to a professionally formatted Word document."""
    if output_dir is None:
        output_dir = _ensure_output_dir(topic)
    docx_path = output_dir / "literature_review.docx"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        if level == 1:
            heading_style.font.size = Pt(18)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    title = doc.add_heading(f"Literature Review: {topic}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()

    _md_to_docx(doc, summary_md)

    doc.save(str(docx_path))
    _log(f"DOCX saved to {docx_path}")
    return docx_path


def _md_to_docx(doc: Document, md_text: str):
    """Parse markdown and add formatted content to the Word document."""
    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("|") and stripped.endswith("|"):
            if _is_separator_row(stripped):
                i += 1
                continue
            row_cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(row_cells)
            if i + 1 < len(lines) and not lines[i + 1].strip().startswith("|"):
                _add_table(doc, table_rows)
                table_rows = []
            elif i + 1 >= len(lines):
                _add_table(doc, table_rows)
                table_rows = []
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        elif stripped.startswith("---") or stripped.startswith("***"):
            pass
        elif stripped == "":
            pass
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

        i += 1


def _is_separator_row(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:]+(\|[\s\-:]+)*\|$", line))


def _add_table(doc: Document, rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

    doc.add_paragraph()


def _add_formatted_text(paragraph, text: str):
    """Add text with inline formatting (bold, italic, PMID citations)."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|\[PMID:\s*\d+\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif re.match(r"\[PMID:\s*\d+\]", part):
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.italic = True
        else:
            paragraph.add_run(part)


def generate_outputs(summary_md: str, topic: str) -> dict[str, Path]:
    """Generate both .md and .docx output files. Returns paths dict."""
    output_dir = _ensure_output_dir(topic)
    md_path = save_markdown(summary_md, topic, output_dir)
    docx_path = save_docx(summary_md, topic, output_dir)
    return {"md_path": md_path, "docx_path": docx_path, "output_dir": output_dir}
